import os
import streamlit as st
import fitz  # PyMuPDF
import openai
from pinecone import Pinecone
import json
import re
from docx import Document
from io import BytesIO

# Initialize API Keys
pinecone_key = os.environ.get("PINECONE_API_KEY")
openai_key = os.environ.get("OPENAI_API_KEY")

if not pinecone_key or not openai_key:
    st.error("Missing API keys. Please set PINECONE_API_KEY and OPENAI_API_KEY in your environment.")
    st.stop()

# Pinecone and OpenAI Initialization
pc = Pinecone(api_key=pinecone_key, environment="us-east-1")
index = pc.Index("bid-automation-index")
client = openai.OpenAI(api_key=openai_key)

# Streamlit UI Setup
st.set_page_config(page_title="Bid Proposal Generator", layout="wide")
st.title("📂 Bid Proposal Generator")
# Automated Bid Preparation System (ABPSys)

# Session State Defaults
for key in ["requirements", "similar_proposals", "methodology_text", "detailed_sections"]:
    if key not in st.session_state:
        st.session_state[key] = ""

# Helper: Format Proposal Content
def format_proposal_content(text: str) -> str:
    return re.sub(r'\n{3,}', '\n\n', text.strip())

# Helper: Extract Methodology Sections using GPT only
def extract_methodology_sections(proposal_chunks) -> str:
    full_text = "\n\n".join([chunk["metadata"].get("text", "") for chunk in proposal_chunks]).strip()

    prompt = f"""You are an expert in analyzing technical proposals. From the text below, extract and return only the methodology section. 
The methodology section may be labeled differently, such as 'Approach', 'Work Plan', 'Execution Strategy', or 'Technical Approach', and it typically explains how the project will be executed.

Proposal Text:
{full_text}
"""

    full_response = ""
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "system", "content": prompt}],
        max_tokens=6000  #previously 4k
    )
    full_response += response.choices[0].message.content

    while response.choices[0].finish_reason == "length":
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": "Continue"}],
            max_tokens=6000
        )
        full_response += response.choices[0].message.content

    return full_response.strip()


# Helper: Extract Requirements from RFP
def extract_requirements(text):
    prompt = f"""You are an AI assistant specializing in extracting key information from Request for Proposals (RFPs).

Extract the key requirements and relevant information needed to respond to the RFP based on the content provided.

Give special attention to the **Terms of Reference (TOR)** and provide enough information about the TOR, as it contains critical details about scope, deliverables, and technical expectations which will guide the methodology and response structure.

Structure your output in the following format. If any item is not found, state "Not specified" or "Not available":

1. Project Title  
2. Project Description / Scope of Services / TOR Highlights  
3. Client / Procuring Entity  
4. Submission Deadline  
5. Submission Requirements  
6. Evaluation Criteria  
7. Required Documents  
8. Eligibility or Qualification Requirements  
9. Country of RFP / Country of Project Execution  
10. Any Other Important Information

RFP Content:
{text}
"""

    full_response = ""
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "system", "content": prompt}],
        max_tokens=4000
    )
    full_response += response.choices[0].message.content

    while response.choices[0].finish_reason == "length":
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": "Continue"}],
            max_tokens=4000
        )
        full_response += response.choices[0].message.content

    try:
        extracted = json.loads(full_response)
        st.session_state.requirements = json.dumps(extracted, indent=2)
        return st.session_state.requirements
    except json.JSONDecodeError:
        st.session_state.requirements = full_response
        return st.session_state.requirements


# File Upload
uploaded_file = st.file_uploader("Upload an RFP", type=["pdf"])

if uploaded_file and not st.session_state.requirements:
    with st.spinner("🔍 Extracting RFP Content..."):
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        text = "\n".join([page.get_text("text") for page in doc])

    with st.spinner("🤖 Extracting requirements from RFP..."):
        requirements = extract_requirements(text)
        st.session_state.requirements = requirements

# Show Extracted Requirements
def strip_markdown(text):
    return re.sub(r"[*_`#>\[\]]+", "", text)
buffer_req = None

requirements = st.session_state.requirements

if requirements.strip():
    clean_requirements = strip_markdown(requirements)

    st.subheader("📄 Bid Requirements")
    st.text_area("Requirements", clean_requirements, height=200)

    doc_req = Document()
    doc_req.add_heading("Bid Requirements", level=1)
    doc_req.add_paragraph(clean_requirements)  # Use cleaned text here

    buffer_req = BytesIO()
    doc_req.save(buffer_req)
    buffer_req.seek(0)

# Extract search query fields
def extract_search_query_fields(requirements_json: str) -> str:
    """
    Extracts and formats relevant fields from the RFP requirements
    for use in searching similar proposals.
    """
    try:
        reqs = json.loads(requirements_json)
    except json.JSONDecodeError:
        return requirements_json  # Fallback to raw input if not valid JSON

    fields_to_use = [
        "Project Title",
        "Project Description",
        "Scope of Services",
        "TOR Highlights",
        "Evaluation Criteria",
        "Eligibility or Qualification Requirements",
        "Country of RFP",
        "Country of Project Execution"
    ]

    extracted_fields = []
    for field in fields_to_use:
        if field in reqs and reqs[field].strip():
            extracted_fields.append(f"{field}: {reqs[field].strip()}")

    return "\n".join(extracted_fields)


# Create two columns for side-by-side buttons
if "requirements" in st.session_state and st.session_state.requirements.strip():
    col1, col2 = st.columns([1, 1])

    with col1:
        if st.button("🔎 Get Similar Proposals"):
            with st.spinner("🔍 Searching in Database..."):
                search_input = extract_search_query_fields(requirements)

                embedding_response = client.embeddings.create(
                    input=search_input,
                    model="text-embedding-ada-002"
                )

                # embedding_response = client.embeddings.create(
                #     input=requirements,
                #     model="text-embedding-ada-002"
                # )

                query_embedding = embedding_response.data[0].embedding

                results = index.query(
                    vector=query_embedding,
                    top_k=15,
                    include_metadata=True,
                    filter={"doc_type": {"$eq": "technical"}}
                )

                source_to_matches = {}
                for match in results["matches"]:
                    source = match["metadata"].get("source")
                    if source:
                        source_to_matches.setdefault(source, []).append(match)

                source_scores = []
                for source, matches in source_to_matches.items():
                    avg_score = sum([m["score"] for m in matches]) / len(matches)
                    source_scores.append((source, avg_score, matches))

                top_sources = sorted(source_scores, key=lambda x: x[1], reverse=True)[:3]
                if not top_sources:
                    st.info("No similar proposals found.")
                    st.stop()

                st.session_state.similar_proposals = []

                for i, (source, avg_score, _) in enumerate(top_sources, start=1):
                    all_chunks = index.query(
                        vector=query_embedding,
                        top_k=100,
                        include_metadata=True,
                        filter={"source": {"$eq": source}}
                    )

                    sorted_chunks = sorted(
                        all_chunks["matches"],
                        key=lambda x: x["metadata"].get("chunk_index", 0)
                    )

                    full_text = format_proposal_content(
                        "\n\n".join([chunk["metadata"].get("text", "") for chunk in sorted_chunks])
                    )

                    methodology_text = extract_methodology_sections(sorted_chunks)

                    st.session_state.similar_proposals.append({
                        "title": f"Proposal {i} (Source: {source}, Score: {avg_score:.3f})",
                        "content": full_text,
                        "methodology_only": methodology_text
                    })

    with col2:
        if buffer_req:
            st.download_button(
                label="💾 Download Bid Requirements (Word)",
                data=buffer_req,
                file_name="Bid_requirements.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# Show Top 3 Matching Proposals
if st.session_state.similar_proposals:
    st.subheader("📜 Top 3 Matching Proposals")
    for i, prop in enumerate(st.session_state.similar_proposals):
        with st.expander(prop["title"]):
            st.text_area(f"📄 Full Proposal Content", prop["content"], height=300)
            if prop["methodology_only"]:
                st.text_area("📌 Extracted Methodology Section", prop["methodology_only"], height=200)


# Generate New Methodology Statement

# -------- Extract Section Headings --------
def extract_sections(requirements, top_methodologies):
    prompt = f"""
    You are a proposal development expert.

    Your task is to generate appropriate section titles for a consultant's Methodology Statement. Use the following Extracted Requirements only to infer project context, but DO NOT copy section headings directly from them (e.g., 'Evaluation Criteria', 'Eligibility Requirements' etc. are not appropriate methodology sections).

    Instead, derive professional, context-aware Methodology Statement sections that align with both the project requirements and the tone/structure of the Reference Methodologies.

    Guidelines:
    - Return only numbered section titles.
    - Ensure the structure closely resembles the Reference Methodologies.
    - Do not include administrative or procurement-related terms as section headings.
    - The methodology must reflect the consultant's approach, not the client's instructions.

    Extracted Requirements:
    {requirements}

    Reference Methodologies:
    {top_methodologies}
    """
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2,
        max_tokens=1000
    )
    return response.choices[0].message.content.split('\n')

# -------- Generate Section Content --------
def generate_section_content(section_title, requirements, top_methodologies):
    prompt = f"""
    Write a full and detailed Methodology Statement section titled "{section_title}" for a consulting assignment in infrastructure, based on the following:

    Requirements:
    {requirements}

    Reference Methodologies:
    {top_methodologies}

    Instructions:
    - Provide a deep, professional write-up (no placeholders).
    - Limit the total output to what would reasonably fit within a 20-page Word document, assuming standard margins and font sizes.
    - Prioritize clarity, precision, and essential technical detail over exhaustive elaboration.
    - Avoid unnecessary repetition or generic phrasing.
    - Use paragraphs to develop sections as appropriate, only add subheadings when necessary.
    - Highlight relevant technical expertise, tools, or innovative approaches where applicable.
    - Reflect the tone, style, and structure of the reference methodologies (including country- or sector-specific nuances).
    - Prioritize alignment with the client’s needs, objectives, and expectations.
    - Do not include a conclusion for any of the sections, subsections, or for the methodology as a whole or even include it as a section on it's own, unless explicitly stated in the requirements.
    """

    # - Provide a deep, professional write-up (no placeholders).
    # - Include tools, techniques, timeframes, personnel, and approaches.
    # - Follow the inferred tone and style of the reference methodologies.
    # - Do not write a generic intro or conclusion.

    # - Reflect the specific client requirements and project context.
    # - Match the style and tone of the reference methodologies (including country- or sector-specific nuances).
    # - Provide clear, logical steps and justifications for major proposed activities.
    # - Highlight relevant technical expertise, tools, or innovative approaches where applicable.
    # - Avoid vague or overly generic statements. Be specific and evidence-based.
    # - Do not follow a rigid structure — let the format emerge naturally based on the content.
    # - Do not include a generic Conclusion section unless explicitly required in the extracted requirements.
    # - 

    # """
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4,
        max_tokens=6000
    )
    return response.choices[0].message.content

if st.session_state.similar_proposals:
    if st.button("📝 Generate Methodology and Work Plan"):
        with st.spinner("🤖 Generating Methodology and Work Plan..."):
            top_methodologies = "\n\n".join([
                p["methodology_only"] for p in st.session_state.similar_proposals if p["methodology_only"]
            ])

            # Step 1: Extract inferred section titles
            sections = extract_sections(requirements, top_methodologies)

            # Initialize the progress bar
            progress_bar = st.progress(0)
            total_sections = len([s for s in sections if s.strip()])
            section_counter = 0

            # Step 2: Generate content for each section
            full_methodology = ""
            for section in sections:
                if section.strip():
                    with st.spinner(f"✍️ Writing section: {section.strip()}"):
                        content = generate_section_content(section.strip(), requirements, top_methodologies)
                        full_methodology += f"\n\n{content.strip()}"

                    # Update progress
                    section_counter += 1
                    progress_percent = int((section_counter / total_sections) * 100)
                    progress_bar.progress(progress_percent)

            # Ensure bar ends at 100% even if rounding missed it
            progress_bar.progress(100)
            st.session_state.methodology_text = full_methodology


# Show and Download Methodology
def strip_markdown(text):
    return re.sub(r"[*_`#>\[\]]+", "", text)

def clean_markdown(text):
    text = re.sub(r"[ \t]+", " ", text)  # Collapse multiple spaces
    text = re.sub(r"\r\n|\r", "\n", text)  # Normalize line endings
    text = re.sub(r"[*_`#>\[\]]+", "", text)  # Remove Markdown/special characters
    return text.strip()


if st.session_state.methodology_text:
    clean_text = strip_markdown(st.session_state.methodology_text)
    st.subheader("📌 Methodology and Work Plan")
    st.text_area("Generated Methodology Section", clean_text, height=300)


    doc = Document()
    doc.add_heading("Methodology and Work Plan", level=1)
    doc.add_paragraph(clean_text)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    
    st.download_button(
            label="💾 Download Methodology and Work Plan (Word)",
            data=buffer,
            file_name="Methodology_and_work_plan.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )