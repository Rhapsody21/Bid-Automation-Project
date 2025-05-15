import streamlit as st
import os
import openai
import fitz  # PyMuPDF
import tempfile
from dotenv import load_dotenv
import re
import random
import string
import json
from docx import Document
from io import BytesIO
import glob

# Import required libraries
from pinecone import Pinecone as PineconeClient
from langchain_openai import OpenAIEmbeddings
from langchain.llms import OpenAI

# Load environment variables
load_dotenv()

# Configuration and API Key Setup
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
INDEX_NAME = os.getenv("PINECONE_INDEX_NAME", "bid-project-index")

# Initialize Clients
try:
    pc_client = PineconeClient(api_key=PINECONE_API_KEY)
    index = pc_client.Index(name=INDEX_NAME)
    embeddings = OpenAIEmbeddings(api_key=OPENAI_API_KEY)
except Exception as e:
    st.error(f"Error initializing clients: {str(e)}")

# Streamlit Page Configuration
st.set_page_config(page_title="RFP Proposal Generator", layout="wide")

class ProposalGenerator:
    @staticmethod
    def fullscreen_text_area(text, height=300, key=None):
        """
        Create a fullscreen text viewer with a floating button.
        
        Args:
            text (str): Text to display
            height (int): Initial height of text area
            key (str, optional): Unique key for the component
        
        Returns:
            str: Displayed text
        """
        if key is None:
            key = ''.join(random.choices(string.ascii_lowercase, k=10))
        
        # Regular text area for input/output
        text_area = st.text_area(f"Content", text, height=height, key=key)
        
        # Simplified HTML for fullscreen button
        html = f"""
        <div style="position: relative; width: 100%; margin-top: -45px; z-index: 1000;">
            <button onclick="toggleFullScreen('{key}')" style="
                position: absolute;
                bottom: 20px;
                right: 20px;
                width: 40px;
                height: 40px;
                border-radius: 50%;
                background-color: rgba(0, 123, 255, 0.7);
                color: white;
                border: none;
                cursor: pointer;
                z-index: 1000;">
                <svg xmlns="http://www.w3.org/2000/svg" width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2">
                    <path d="M8 3H5a2 2 0 0 0-2 2v3m18 0V5a2 2 0 0 0-2-2h-3m0 18h3a2 2 0 0 0 2-2v-3M3 16v3a2 2 0 0 0 2 2h3"></path>
                </svg>
            </button>
        </div>"""
        st.components.v1.html(html, height=0)
        
        return text_area

    @staticmethod
    def extract_text_from_pdf(pdf_path):
        """
        Extract text from a PDF file.
        
        Args:
            pdf_path (str): Path to the PDF file
        
        Returns:
            str: Extracted text from PDF
        """
        try:
            doc = fitz.open(pdf_path)
            text = ""
            for page in doc:
                text += page.get_text("text") + "\n\n"
            doc.close()  # Properly close the document
            return text
        except Exception as e:
            st.error(f"Error extracting PDF text: {e}")
            return ""

    @staticmethod
    def extract_key_words(rfp_content):
        """
        Extract key words from RFP content using OpenAI.
        
        Args:
            rfp_content (str): RFP document content
        
        Returns:
            str: Extracted key phrases
        """
        try:
            openai.api_key = OPENAI_API_KEY

            # Truncate content if too long
            max_length = 15000
            truncated_content = rfp_content[:max_length] if len(rfp_content) > max_length else rfp_content

            response = openai.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert RFP analysis assistant. Extract critical keywords and phrases that define the proposal's core requirements."
                    },
                    {
                        "role": "user",
                        "content": truncated_content
                    }
                ],
                max_tokens=500
            )
            return response.choices[0].message.content
        except Exception as e:
            st.error(f"Error extracting key words: {e}")
            return "Error extracting key requirements. Please try again."

    @staticmethod
    def search_similar_proposals(query):
        """
        Search for similar technical proposals.
        
        Args:
            query (str): Search query
        
        Returns:
            list: Top matching proposals
        """
        try:
            query_embedding = embeddings.embed_query(query)
            
            results = index.query(
                vector=query_embedding,
                filter={"collection_type": "technical_proposal"},
                top_k=3,
                include_metadata=True
            )

            matches = results.matches if hasattr(results, 'matches') else results.get("matches", [])
            
            if not matches:
                return "No similar proposals found."

            matched_proposals = []
            for match in matches:
                metadata = match.metadata if hasattr(match, 'metadata') else match.get("metadata", {})
                filename = metadata.get("filename", "Unknown")
                score = match.score if hasattr(match, 'score') else match.get("score", 0)
                
                try:
                    # Fetch full document chunks with updated query format
                    chunks_query = index.query(
                        vector=query_embedding,
                        filter={"collection_type": "technical_proposal", "filename": filename},
                        top_k=50,
                        include_metadata=True
                    )
                    
                    chunk_matches = chunks_query.matches if hasattr(chunks_query, 'matches') else chunks_query.get("matches", [])
                    
                    document_chunks = []
                    for chunk in chunk_matches:
                        chunk_metadata = chunk.metadata if hasattr(chunk, 'metadata') else chunk.get("metadata", {})
                        document_chunks.append({
                            "chunk_id": chunk_metadata.get("chunk_id", 0),
                            "text": chunk_metadata.get("text", "")
                        })
                    
                    # Sort chunks by ID to maintain document order
                    document_chunks.sort(key=lambda x: x["chunk_id"])
                    complete_text = "\n\n".join([chunk["text"] for chunk in document_chunks])
                    matched_proposals.append((filename, score, complete_text))
                except Exception as chunk_error:
                    st.warning(f"Error retrieving chunks for {filename}: {str(chunk_error)}")
                    # Still add the proposal with partial content
                    matched_proposals.append((filename, score, f"Error retrieving full content: {str(chunk_error)}"))
            
            return matched_proposals
        except Exception as e:
            st.error(f"Error searching similar proposals: {str(e)}")
            return []

    @staticmethod
    def generate_proposal(rfp_content, key_phrases, similar_proposals):
        """
        Generate a comprehensive proposal using OpenAI.
        
        Args:
            rfp_content (str): RFP document content
            key_phrases (str): Extracted key phrases
            similar_proposals (list): Similar proposal documents
        
        Returns:
            str: Generated proposal
        """
        try:
            openai.api_key = OPENAI_API_KEY

            # Prepare reference proposals with length limit to avoid token limits
            reference_content = ""
            for _, score, content in similar_proposals[:3]:
                truncated_content = content[:5000]  # Limit each proposal to 5000 chars
                reference_content += f"\nEXAMPLE PROPOSAL (Score: {score:.2f}):\n{truncated_content}\n\n"

            # Truncate RFP content if needed
            max_rfp_length = 10000
            truncated_rfp = rfp_content[:max_rfp_length] if len(rfp_content) > max_rfp_length else rfp_content

            response = openai.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": "You are a senior proposal writer specializing in crafting winning technical proposals."
                    },
                    {
                        "role": "user",
                        "content": f"""
                        Generate a comprehensive technical proposal methodology based on:
                        
                        RFP KEY REQUIREMENTS:
                        {key_phrases}
                        
                        REFERENCE PROPOSALS:
                        {reference_content}
                        
                        Guidelines:
                        - Create a detailed, structured methodology statement
                        - Use professional, technical language
                        - Demonstrate deep understanding of project requirements
                        - Incorporate best practices from reference proposals
                        - For each heading make sure you use ### heading signs.
                        """
                    }
                ],
                max_tokens=16000,
                temperature=0.7
            )

            return response.choices[0].message.content
        except Exception as e:
            st.error(f"Proposal generation error: {e}")
            return f"Error generating proposal: {str(e)}"

    @staticmethod
    def extract_headings(text):
        """
        Extract markdown headings from text.
        
        Args:
            text (str): Input text
        
        Returns:
            list: List of headings
        """
        if not text:
            return []
            
        pattern = r"^(#{1,4}\s.*?)$"
        return re.findall(pattern, text, re.MULTILINE)
    
    @staticmethod
    def extract_section_content(generated_proposal, selected_section):
        """
        Extract the content of a specific section from the generated proposal.
        
        Args:
            generated_proposal (str): Full generated proposal text
            selected_section (str): Section heading to extract
        
        Returns:
            str: Content of the specified section
        """
        if not generated_proposal or not selected_section:
            return "No proposal or section selected."
            
        # Escape special regex characters in the section name
        escaped_section = re.escape(selected_section.strip())
        
        # Regex pattern to find section content
        # This will capture all text from the section heading to the next heading or end of document
        pattern = rf"{escaped_section}(.*?)(?=\n#|\Z)"
        
        # Search for the section content
        match = re.search(pattern, generated_proposal, re.DOTALL | re.MULTILINE | re.IGNORECASE)
        
        if match:
            return match.group(1).strip()
        else:
            return "Section not found in the proposal."
    
    @staticmethod
    def update_section_in_proposal(full_proposal, section_heading, new_section_content):
        """
        Update a specific section in the proposal with new content.
        
        Args:
            full_proposal (str): Complete proposal text
            section_heading (str): Heading of the section to update
            new_section_content (str): New content for the section
        
        Returns:
            str: Updated proposal with the modified section
        """
        if not full_proposal or not section_heading:
            return full_proposal
            
        # Escape special regex characters in the section name
        escaped_section = re.escape(section_heading.strip())
        
        # Find the section and its content up to the next heading or end of document
        pattern = rf"({escaped_section})(.*?)(?=\n#|\Z)"
        
        # Search for the pattern in the proposal
        match = re.search(pattern, full_proposal, re.DOTALL | re.MULTILINE | re.IGNORECASE)
        
        if match:
            # Create the updated section with original heading and new content
            updated_section = f"{section_heading}\n{new_section_content}"
            
            # Replace the old section with the updated one
            updated_proposal = full_proposal[:match.start()] + updated_section + full_proposal[match.end():]
            return updated_proposal
        else:
            return full_proposal  # Return original if section not found
    
    @staticmethod
    def extract_methodology_and_after(text, api_key=OPENAI_API_KEY):
        """
        Extract methodology section from text using OpenAI.
        
        Args:
            text (str): Input text
            api_key (str): OpenAI API key
        
        Returns:
            str: Extracted methodology section
        """
        if not text:
            return "No text provided for extraction."
            
        try:
            openai.api_key = api_key
            
            # Truncate text if too long
            max_length = 15000
            truncated_text = text[:max_length] if len(text) > max_length else text
            
            response = openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a document parsing expert."},
                    {"role": "user", "content": (
                        "You are a helpful assistant. Given the following academic or report-like document, "
                        "extract the full 'Methodology' section and everything that follows it. "
                        "Do not stop until the next major section such as 'Results', 'Discussion', or 'Conclusion'.\n\n"
                        "DOCUMENT:\n" + truncated_text + "\n\n"
                        "###\n\n"
                        "Return the whole content of the Methodology section."
                    )}
                ],
                temperature=0.2,
                max_tokens=2000
            )
            
            methodology_section = response.choices[0].message.content.strip()
            return methodology_section

        except Exception as e:
            print(f"Error extracting methodology: {e}")
            return f"Error extracting methodology: {str(e)}"
            
    @staticmethod
    def expand_section(section_content):
        """
        Expand a section using OpenAI to provide more detailed explanation.
        
        Args:
            section_content (str): Content of the section to expand
        
        Returns:
            str: Expanded and detailed explanation of the section
        """
        if not section_content:
            return "No section content provided for expansion."
            
        try:
            openai.api_key = OPENAI_API_KEY
            response = openai.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert technical writer specializing in providing comprehensive, detailed explanations of sections of methodology statement."
                    },
                    {
                        "role": "user",
                        "content": f"""
                        Provide an extremely detailed, comprehensive expansion of the following section. 
                        Go beyond the original text. 
                        - First of all make sure you understand well the aim of that section
                        - Compare the aim with what was previously provided as the section content 
                        - Now explain in more details what was given make sure that everything is clearly understandable
                        and detailed as enough as possible. Keep the section title as the section heading because this expanded version will later be incorporated in the 
                        full methodology statement. 

                        Original Section Content:
                        {section_content}

                        Expansion Guidelines:
                        - Use professional, technical language
                        - Provide clear, structured explanations
                        - Add significant depth and context
                        - Demonstrate expert-level understanding
                        """
                    }
                ],
                max_tokens=2000,
                temperature=0.8
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Error in section expansion: {str(e)}"
            
    @staticmethod
    def extract_sections(generated_proposal, api_key=OPENAI_API_KEY):
        """
        Use GPT-4 to extract all sections from the proposal
        
        Args:
            generated_proposal (str): The full proposal text
            api_key (str): OpenAI API key
            
        Returns:
            list: List of section dictionaries with headings and content
        """
        if not generated_proposal:
            return []
            
        try:
            openai.api_key = api_key
            
            # Truncate if too long
            max_length = 20000
            truncated_proposal = generated_proposal[:max_length] if len(generated_proposal) > max_length else generated_proposal
            
            response = openai.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a document parsing expert."},
                    {"role": "user", "content": (
                        "Extract all sections from the following proposal. For each section, provide:\n"
                        "1. The exact section heading (with any markdown formatting like ###)\n"
                        "2. The complete section content\n\n"
                        "Format your response as JSON with this structure:\n"
                        "{\n"
                        "  \"sections\": [\n"
                        "    {\n"
                        "      \"heading\": \"### Section Title\",\n"
                        "      \"content\": \"Full section content goes here...\"\n"
                        "    },\n"
                        "    ...\n"
                        "  ]\n"
                        "}\n\n"
                        "PROPOSAL:\n" + truncated_proposal
                    )}
                ],
                response_format={"type": "json_object"},
                temperature=0.2,
                max_tokens=8000
            )
            
            result = json.loads(response.choices[0].message.content)
            return result.get("sections", [])
        except Exception as e:
            print(f"Error extracting sections with GPT-4: {e}")
            return []

    @staticmethod
    def save_sections_to_files(sections, output_directory="proposal_sections"):
        """
        Save each section to a separate file.
        
        Args:
            sections (list): List of section dictionaries
            output_directory (str): Directory to save files
        
        Returns:
            list: List of saved filenames
        """
        # Delete the previous files in the directory
        folder_path = output_directory # Replace with the full path if needed

        # Get all .txt files in the folder
        txt_files = glob.glob(os.path.join(folder_path, "*.txt"))

        # Delete each file
        for file_path in txt_files:
            try:
                os.remove(file_path)
                print(f"Deleted previous .txt files from the file: {file_path}")
            except Exception as e:
                print(f"Failed to delete {file_path}: {e}")
        
        if not sections:
            return []
        
        # Create directory if it doesn't exist
        try:
            os.makedirs(output_directory, exist_ok=True)
        except Exception as e:
            print(f"Error creating directory: {e}")
            return []
        
        saved_files = []
        
        for i, section in enumerate(sections):
            heading = section.get("heading", "")
            content = section.get("content", "")
            
            if not heading:
                continue
                
            # Clean heading for filename
            clean_heading = re.sub(r'[^a-zA-Z0-9 ]', '', heading.replace('#', '').strip())
            filename = f"{i+1:02d}_{clean_heading.lower().replace(' ', '_')}.txt"
            
            file_path = os.path.join(output_directory, filename)
            
            try:
                with open(file_path, "w", encoding="utf-8") as file:
                    file.write(f"{heading}\n\n{content}")
                saved_files.append(filename)
            except Exception as e:
                print(f"Error saving section '{heading}' to file: {e}")
        
        return saved_files

    @staticmethod
    def update_section_file(heading, content, output_directory="proposal_sections"):
        """
        Update a specific section file with new content.
        
        Args:
            heading (str): Heading of the section to update
            content (str): New content for the section
            output_directory (str): Directory containing section files
        
        Returns:
            bool: True if successful, False otherwise
        """
        if not heading or not content:
            return False
            
        try:
            # Clean heading for filename search
            clean_heading = re.sub(r'[^a-zA-Z0-9 ]', '', heading.replace('#', '').strip())
            search_term = clean_heading.lower().replace(' ', '_')
            
            # Find matching file in directory
            if not os.path.exists(output_directory):
                os.makedirs(output_directory, exist_ok=True)
                
            matching_files = [f for f in os.listdir(output_directory) 
                             if f.endswith('.txt') and search_term in f.lower()]
            
            if not matching_files:
                # Create new file if not found
                filename = f"{clean_heading.lower().replace(' ', '_')}.txt"
                file_path = os.path.join(output_directory, filename)
            else:
                file_path = os.path.join(output_directory, matching_files[0])
            
            # Update file with new content
            with open(file_path, "w", encoding="utf-8") as file:
                file.write(f"{heading}\n\n{content}")
            
            return True
        except Exception as e:
            print(f"Error updating section file: {e}")
            return False
    
    @staticmethod
    def save_proposal_to_file(proposal_text, filename="generated_proposal.txt"):
        """
        Save the generated proposal to a file.
        
        Args:
            proposal_text (str): The proposal text to save
            filename (str): Name of the file to save
        
        Returns:
            bool: True if successful, False otherwise
        """
        if not proposal_text:
            return False
            
        try:
            with open(filename, "w", encoding="utf-8") as file:
                file.write(proposal_text)
            return True
        except Exception as e:
            print(f"Error saving proposal to file: {e}")
            return False
    
    @staticmethod
    def merge_sections_from_files(directory="proposal_sections"):
        """
        Merge all section files into a complete proposal.
        
        Args:
            directory (str): Directory containing section files
        
        Returns:
            str: Merged proposal text
        """
        if not os.path.exists(directory):
            return ""
            
        try:
            # Get all txt files in the directory
            files = [f for f in os.listdir(directory) if f.endswith('.txt')]
            
            # Sort files by their numeric prefix if present
            def get_file_order(filename):
                match = re.match(r'^(\d+)_', filename)
                return int(match.group(1)) if match else float('inf')
                
            files.sort(key=get_file_order)
            
            # Merge content from all files
            merged_content = ""
            for filename in files:
                file_path = os.path.join(directory, filename)
                try:
                    with open(file_path, "r", encoding="utf-8") as file:
                        merged_content += file.read() + "\n\n"
                except Exception as e:
                    print(f"Error reading file {filename}: {e}")
            
            return merged_content
        except Exception as e:
            print(f"Error merging sections: {e}")
            return ""

def main():
    st.title("Technical Proposal Generator")
    
    # Initialize session state variables if they don't exist
    session_vars = [
        'uploaded_file', 'rfp_text', 'key_phrases', 'similar_proposals', 
        'generated_proposal', 'proposal_display', 'sections', 
        'section_content', 'expanded_section', 'displayed_expanded_content',
        'displayed_section_content', 'selected_section_title', 
        'section_updated', 'edit_expanded_section', 'extracted_sections',
        'saved_section_files'
    ]
    
    for key in session_vars:
        if key not in st.session_state:
            st.session_state[key] = None

    # Callbacks for buttons to avoid rerunning on state change
    def generate_proposal_callback():
        st.session_state.generate_proposal_clicked = True
    
    def view_section_callback():
        st.session_state.view_section_clicked = True
    
    def expand_section_callback():
        st.session_state.expand_section_clicked = True
    
    def update_section_callback():
        # Store the updated content in a temporary session state variable
        st.session_state.temp_updated_content = st.session_state.edit_expanded_section
        st.session_state.update_section_clicked = True
    
    def discard_update_callback():
        st.session_state.discard_update_clicked = True
    
    def extract_sections_callback():
        st.session_state.extract_sections_clicked = True
    
    # File uploader
    uploaded_file = st.file_uploader("Upload RFP Document", type=["pdf"])
    
    # Process new file upload
    if uploaded_file and (st.session_state.uploaded_file is None or 
                         (hasattr(uploaded_file, 'name') and 
                          hasattr(st.session_state.uploaded_file, 'name') and 
                          uploaded_file.name != st.session_state.uploaded_file.name)):
        st.session_state.uploaded_file = uploaded_file
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                temp_file.write(uploaded_file.read())
                temp_file_path = temp_file.name

            st.success("File uploaded successfully")
            
            # Extract and display RFP content
            rfp_text = ProposalGenerator.extract_text_from_pdf(temp_file_path)
            
            # Clean up the temporary file
            try:
                os.unlink(temp_file_path)
            except Exception as e:
                print(f"Error removing temporary file: {e}")
                
            if not rfp_text:
                st.error("Could not extract text from the PDF. Please check if the file is valid and not encrypted.")
                return
                
            st.session_state.rfp_text = rfp_text
            
            # Reset related state variables on new file upload
            for key in session_vars[2:]:  # Skip uploaded_file and rfp_text
                st.session_state[key] = None
                
            # Initialize sections as empty list instead of None
            st.session_state.sections = []
            st.session_state.extracted_sections = []
            st.session_state.saved_section_files = []
        except Exception as e:
            st.error(f"Error processing uploaded file: {str(e)}")
            return
        
    # Display RFP content if available
    if st.session_state.rfp_text:
        st.subheader("RFP Content")
        with st.expander("View RFP content", expanded=False):
            ProposalGenerator.fullscreen_text_area(st.session_state.rfp_text, height=300, key="rfp_content")
        
        # Extract key phrases if not already done
        if st.session_state.key_phrases is None:
            with st.spinner("Analyzing RFP..."):
                key_phrases = ProposalGenerator.extract_key_words(st.session_state.rfp_text)
                st.session_state.key_phrases = key_phrases
        
        # Display key phrases if available
        if st.session_state.key_phrases:
            st.subheader("Extracted Key Requirements")
            with st.expander("View extracted key requirements", expanded=False):
                ProposalGenerator.fullscreen_text_area(st.session_state.key_phrases, height=200, key="key_phrases")
            
            # Search for similar proposals if not already done
            if st.session_state.similar_proposals is None:
                with st.spinner("Finding similar proposals..."):
                    similar_proposals = ProposalGenerator.search_similar_proposals(st.session_state.key_phrases)
                    st.session_state.similar_proposals = similar_proposals
            
            # Display similar proposals if available
            if isinstance(st.session_state.similar_proposals, list) and st.session_state.similar_proposals:
                st.subheader("Similar Technical Proposals")
                for idx, (filename, score, content) in enumerate(st.session_state.similar_proposals, 1):
                    with st.expander(f"Proposal {idx}: {filename} (Similarity: {score:.2f})", expanded=False):
                        ProposalGenerator.fullscreen_text_area(content, height=500, key=f"proposal_{idx}")
                    
                    # Extract methodology for each proposal
                    with st.spinner(f"Extracting methodology for proposal {idx}..."):
                        extracted_methodology = ProposalGenerator.extract_methodology_and_after(content)
                        with st.expander(f"Extracted methodology for proposal {idx}: {filename}", expanded=False):
                            ProposalGenerator.fullscreen_text_area(extracted_methodology, height=500, key=f"methodology_{idx}")
                
                # Generate proposal button
                if st.button("Generate Proposal", on_click=generate_proposal_callback):
                    pass
                
                # Handle proposal generation
                if 'generate_proposal_clicked' in st.session_state and st.session_state.generate_proposal_clicked:
                    if st.session_state.generated_proposal is None:
                        with st.spinner("Generating comprehensive proposal..."):
                            try:
                                generated_proposal = ProposalGenerator.generate_proposal(
                                    st.session_state.rfp_text, 
                                    st.session_state.key_phrases, 
                                    st.session_state.similar_proposals
                                )
                                
                                if not generated_proposal or "Error generating proposal" in generated_proposal:
                                    st.error("Failed to generate proposal. Please try again.")
                                else:
                                    st.session_state.generated_proposal = generated_proposal
                                    st.session_state.proposal_display = generated_proposal  # Add display copy
                                    st.session_state.sections = ProposalGenerator.extract_headings(generated_proposal)
                                    
                                    # Save the initial proposal to file
                                    ProposalGenerator.save_proposal_to_file(generated_proposal)
                                    
                                    # Extract sections using GPT-4
                                    with st.spinner("Using GPT-4 to extract and save proposal sections..."):
                                        extracted_sections = ProposalGenerator.extract_sections(generated_proposal)
                                        st.session_state.extracted_sections = extracted_sections
                                        
                                        # Save extracted sections to files
                                        output_dir = os.path.join(os.getcwd(), "proposal_sections")
                                        saved_files = ProposalGenerator.save_sections_to_files(
                                            extracted_sections, 
                                            output_directory=output_dir
                                        )
                                        st.session_state.saved_section_files = saved_files
                                        
                                        # Display success message with saved files
                                        if saved_files:
                                            st.success(f"Successfully saved {len(saved_files)} sections to '{output_dir}' directory")
                                            with st.expander("Saved section files", expanded=True):
                                                for file in saved_files:
                                                    st.write(f"- {file}")
                                        else:
                                            st.warning("No sections were saved. Please check if the proposal contains properly formatted sections.")
                            except Exception as e:
                                st.error(f"Error during proposal generation: {str(e)}")
                    
                    # Reset the flag
                    st.session_state.generate_proposal_clicked = False
            elif st.session_state.similar_proposals == "No similar proposals found.":
                st.warning("No similar proposals found in the database. The generated proposal may be less tailored.")
                # Still allow generating a proposal
                if st.button("Generate Proposal Anyway", on_click=generate_proposal_callback):
                    pass
    
    # Display generated proposal and section tools if available
    if st.session_state.generated_proposal:
        st.subheader("Generated Proposal")
        with st.expander("View full generated proposal", expanded=False):
            proposal_display = st.session_state.proposal_display if st.session_state.proposal_display else st.session_state.generated_proposal
            ProposalGenerator.fullscreen_text_area(proposal_display, height=600, key="proposal_display_view")
        
        # Create Word document for download
        try:
            doc = Document()
            doc.add_heading("Generated Technical Proposal", level=1)
            
            # Split the proposal by headings and add them properly to the document
            lines = proposal_display.split('\n')
            current_text = ""
            
            for line in lines:
                if line.strip().startswith('#'):
                    # Add accumulated text before adding new heading
                    if current_text:
                        doc.add_paragraph(current_text)
                        current_text = ""
                    
                    # Add heading with appropriate level
                    heading_level = line.count('#')
                    heading_text = line.replace('#', '').strip()
                    doc.add_heading(heading_text, level=min(heading_level, 9))
                else:
                    current_text += line + '\n'
            
            # Add any remaining text
            if current_text:
                doc.add_paragraph(current_text)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Download button for Word doc
            st.download_button(
                label="💾 Download Proposal as Word Document",
                data=buffer,
                file_name="generated_proposal.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Error creating Word document: {str(e)}")
            
            # Fallback to plain text download
            st.download_button(
                label="💾 Download Proposal as Text",
                data=proposal_display,
                file_name="generated_proposal.txt",
                mime="text/plain"
            )
        
        # Section selection and operations
        st.subheader("Proposal Sections")
        
        # Section selector dropdown
        if st.session_state.sections:
            selected_section = st.selectbox(
                "Choose a section:", 
                st.session_state.sections,
                key="section_selector"
            )
            st.session_state.selected_section_title = selected_section
            
            # Buttons for section operations in columns
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("View Selected Section", on_click=view_section_callback):
                    pass
            
            with col2:
                if st.button("Expand Selected Section", on_click=expand_section_callback):
                    pass
            
            # Handle view section button
            if 'view_section_clicked' in st.session_state and st.session_state.view_section_clicked:
                with st.spinner(f"Extracting section content..."):
                    section_content = ProposalGenerator.extract_section_content(
                        st.session_state.generated_proposal, 
                        st.session_state.selected_section_title
                    )
                    st.session_state.section_content = section_content
                    st.session_state.displayed_section_content = True
                    st.session_state.displayed_expanded_content = False
                
                # Reset the flag
                st.session_state.view_section_clicked = False
            
            # Handle expand section button
            if 'expand_section_clicked' in st.session_state and st.session_state.expand_section_clicked:
                with st.spinner(f"Expanding section..."):
                    # First get the section content if not already fetched
                    if not st.session_state.section_content:
                        section_content = ProposalGenerator.extract_section_content(
                            st.session_state.generated_proposal, 
                            st.session_state.selected_section_title
                        )
                        st.session_state.section_content = section_content
                    
                    # Now expand this section
                    expanded_content = ProposalGenerator.expand_section(st.session_state.section_content)
                    st.session_state.expanded_section = expanded_content
                    st.session_state.displayed_expanded_content = True
                    st.session_state.displayed_section_content = False
                    st.session_state.edit_expanded_section = expanded_content  # Store for editing
                
                # Reset the flag
                st.session_state.expand_section_clicked = False

            # Display the section content or expanded content
            if st.session_state.displayed_section_content and st.session_state.section_content:
                st.subheader("Original Section Content")
                ProposalGenerator.fullscreen_text_area(st.session_state.section_content, height=400, key="section_content_view")

            # Display expanded section with edit capability
            if st.session_state.displayed_expanded_content and st.session_state.expanded_section:
                st.subheader("Expanded Section Content")
                st.text_area(
                    "Edit Expanded Content", 
                    value=st.session_state.expanded_section, 
                    height=500,
                    key="edit_expanded_section"
                )
                
                # Add buttons to apply or discard changes
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("Update Section with Expanded Content", on_click=update_section_callback):
                        pass
                with col2:
                    if st.button("Discard Changes", on_click=discard_update_callback):
                        pass

            # Handle update section button
            if 'update_section_clicked' in st.session_state and st.session_state.update_section_clicked:
                with st.spinner("Updating section in the proposal..."):
                    try:
                        # Get the updated content from the session state
                        new_content = st.session_state.edit_expanded_section
                        
                        if new_content:
                            # Update the section in the full proposal
                            updated_proposal = ProposalGenerator.update_section_in_proposal(
                                st.session_state.generated_proposal,
                                st.session_state.selected_section_title,
                                new_content
                            )
                            
                            # Update the proposal display
                            st.session_state.proposal_display = updated_proposal
                            st.session_state.generated_proposal = updated_proposal
                            
                            # Update the section file in the filesystem
                            section_updated = ProposalGenerator.update_section_file(
                                st.session_state.selected_section_title,
                                new_content
                            )
                            
                            if section_updated:
                                st.session_state.section_updated = True
                                st.success(f"Successfully updated section '{st.session_state.selected_section_title}'")
                            else:
                                st.warning("Section file update may have failed, but the proposal was updated in memory.")
                        else:
                            st.warning("No content to update. Please provide some content.")
                    except Exception as e:
                        st.error(f"Error updating section: {str(e)}")
                
                # Reset the flag
                st.session_state.update_section_clicked = False
            
            # Handle discard changes button
            if 'discard_update_clicked' in st.session_state and st.session_state.discard_update_clicked:
                # Reset expanded section editing
                st.session_state.displayed_expanded_content = False
                st.session_state.edit_expanded_section = None
                
                # Reset the flag
                st.session_state.discard_update_clicked = False
                st.warning("Changes discarded")

            # Add a "Regenerate Full Proposal" button to merge all sections
            if st.session_state.saved_section_files:
                if st.button("Regenerate Full Proposal from Section Files"):
                    with st.spinner("Merging sections into complete proposal..."):
                        try:
                            output_dir = os.path.join(os.getcwd(), "proposal_sections")
                            merged_proposal = ProposalGenerator.merge_sections_from_files(directory=output_dir)
                            
                            if merged_proposal:
                                # Update the proposal in the session state
                                st.session_state.generated_proposal = merged_proposal
                                st.session_state.proposal_display = merged_proposal
                                
                                # Save the regenerated proposal
                                saved = ProposalGenerator.save_proposal_to_file(
                                    merged_proposal, 
                                    filename="regenerated_proposal.txt"
                                )
                                
                                if saved:
                                    try:
                                        doc = Document()
                                        doc.add_heading("Methodology statement", level=1)
                                        
                                        # Split the proposal by headings and add them properly to the document
                                        lines = merged_proposal.split('\n')
                                        current_text = ""
                                        
                                        for line in lines:
                                            if line.strip().startswith('#'):
                                                # Add accumulated text before adding new heading
                                                if current_text:
                                                    doc.add_paragraph(current_text)
                                                    current_text = ""
                                                
                                                # Add heading with appropriate level
                                                heading_level = line.count('#')
                                                heading_text = line.replace('#', '').strip()
                                                doc.add_heading(heading_text, level=min(heading_level, 9))
                                            else:
                                                current_text += line + '\n'
                                        
                                        # Add any remaining text
                                        if current_text:
                                            doc.add_paragraph(current_text)

                                        buffer = BytesIO()
                                        doc.save(buffer)
                                        buffer.seek(0)

                                        # Download button for Word doc
                                        st.download_button(
                                            label="💾 Download updated Methodology statement as Word Document",
                                            data=buffer,
                                            file_name="updated_proposal.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                        )
                                    except Exception as e:
                                        st.error(f"Error creating Word document: {str(e)}")
                                        
                                        # Fallback to plain text download
                                        st.download_button(
                                            label="💾 Download updated Methodology statement",
                                            data=merged_proposal,
                                            file_name="regenerated_proposal.txt",
                                            mime="text/plain"
                                        )
                                else:
                                    st.error("Failed to save regenerated proposal")
                            else:
                                st.error("Failed to merge proposal sections")
                        except Exception as e:
                            st.error(f"Error regenerating proposal: {str(e)}")
        else:
            st.warning("No sections found in the generated proposal.")

if __name__ == "__main__":
    main()

